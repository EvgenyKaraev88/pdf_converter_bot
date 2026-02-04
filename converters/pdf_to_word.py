import os
import tempfile
from PyPDF2 import PdfReader
from docx import Document


def convert_pdf_to_word(pdf_path):
    """
    Convert a PDF file to Word (.docx) format.
    This is a basic implementation that extracts text from PDF and creates a Word document.
    
    Args:
        pdf_path (str): Path to the input PDF file
        
    Returns:
        str: Path to the output Word file, or None if conversion fails
    """
    try:
        # Create a temporary Word file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_output:
            output_path = temp_output.name
        
        # Read the PDF file
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PdfReader(pdf_file)
            
            # Create a new Word document
            doc = Document()
            
            # Extract text from each page and add to the Word document
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                # Add the extracted text to the document
                if text.strip():  # Only add if there's text
                    # Split text into paragraphs by double newlines or regular newlines
                    paragraphs = text.split('\n')
                    
                    for paragraph_text in paragraphs:
                        if paragraph_text.strip():  # Only add non-empty paragraphs
                            doc.add_paragraph(paragraph_text.strip())
                        
                        # Add extra space between pages
                        if paragraph_text == paragraphs[-1] and page_num < len(pdf_reader.pages) - 1:
                            doc.add_paragraph()  # Add an empty paragraph as spacing between pages
        
        # Save the Word document
        doc.save(output_path)
        
        return output_path
    
    except Exception as e:
        print(f"Error converting PDF to Word: {str(e)}")
        # Clean up in case of error
        if 'output_path' in locals() and os.path.exists(output_path):
            os.remove(output_path)
        return None